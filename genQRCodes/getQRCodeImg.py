from dataclasses import dataclass
import requests
import configparser
import json
import qrcode
from PIL import Image, ImageDraw, ImageFont
import time
from docx import Document
from docx.shared import Cm, Pt, Inches
import logging
import os
import ast

default_table_url = "https://yesv-desaysv.feishu.cn/base/CmHmb4MxPaEW7zsWB07c1hCUnhd?table=tbl3OBzMMqjX79gN&view=vewsIt61jC#CategoryScheduledTask"
app_id = 'cli_a514aea9fa79900b'
app_secret = 'IsUeIxmzO5NtJiQA6B3MdfkHqIcmQqws'
app_token = 'CmHmb4MxPaEW7zsWB07c1hCUnhd'
table_id = 'tbl3OBzMMqjX79gN'
bg_qrcode = 'white'
bg_label = 'white'
bg_content = 'white'
cl_qrcode = 'black'
cl_text = 'black'

"""log 以追加的形式打印到文件夹./qrcodes下"""
logging.basicConfig(level=logging.INFO, filemode='a', format='%(asctime)s - %(levelname)s - %(message)s')


@dataclass
class TextConf:
    is_show = True
    items: tuple | None = None
    font_size = 30
    left = True


@dataclass
class PageConf:
    row_colum: tuple | None = None

    page_size: tuple | None = None
    page_margin: tuple | None = None

    label_size_cm: tuple | None = None
    label_size_pixels: tuple | None = None
    label_margin_pixels = 30

    content_size_cm: tuple | None = None
    content_size_pixels: tuple | None = None
    content_margin_pixels = 10

    min_direction_pixels = 100

@dataclass
class MobileConf:
    url = default_table_url
    qr_code_data:tuple | None = None
    editable:tuple | None = None

def cm_to_pixels(cm, dpi=300):
    return int(cm * dpi / 2.54)


"""如果文件qrcodes目录不存在，则创建"""
if not os.path.exists('./qrcodes'):
    os.makedirs('./qrcodes')

configure = configparser.ConfigParser()

text_conf = TextConf()
page_conf = PageConf()
mobile_conf = MobileConf()

def check_config_file():
    global configure, cl_qrcode, cl_text
    # 解析并读取配置文件feishu-config.ini
    configure["多维表格地址"] = {'链接地址': default_table_url}
    # 在config文件中加入注释"单位为cm"
    desc = "#以下为默认配置。单位为cm, A4纸打印，大小为21cm*29.7cm。可选常用颜色: white black green red blue yellow"
    configure["配置说明"] = {"说明": desc}
    configure["打印纸张"] = {'边距_上下左右': (1.56, 1.3, 0.78, 0.78)}
    configure["标签"] = {'标签行列': (7, 3)}
    configure["文本显示"] = {
        '字段': ('设备ID', '设备名称', '项目', '录入日期', '使用人'),
        '字体大小': 30,
        '位置_左右': '左'
    }
    configure["颜色"] = {'二维码颜色': 'black', '文本颜色': 'black'}
    configure["二维码内容字段"] = {'字段': ('设备ID', '设备名称', '项目', '录入日期', '使用人')}
    configure["扫码设置"] = {'可编辑字段': ('使用人',)}
    # 检查飞书配置文件  feishu-config.ini，如果不存在，则创建
    if not os.path.exists('./feishu-config.ini'):
        # utf-8 编码写入配置文件
        with open('./feishu-config.ini', 'w', encoding='utf-8') as cf:
            configure.write(cf)
    else:
        configure.read('feishu-config.ini', encoding='utf-8')
    paper_size = (21.0, 29.7)
    configure.set("打印纸张", 'page_size', str(paper_size))
    if not configure.has_section('SECRET'):
        configure.add_section('SECRET')
    configure.set('SECRET', 'app_id', app_id)
    configure.set('SECRET', 'app_secret', app_secret)
    configure.set('SECRET', 'app_token', app_token)
    configure.set('SECRET', 'table_id', table_id)

    # 读取配置文件中的url, 如果没有，则使用默认值
    loc_url = configure.get('多维表格地址', '链接地址', fallback=None)
    logging.info(f'feishu-config.ini url:{loc_url}')
    if not loc_url:
        loc_url = default_table_url
        logging.info(f'feishu-config.ini default url:{loc_url}')

    # 如果url符合格式则从url中解析出app_token table_id
    if loc_url.startswith('https://yesv-desaysv.feishu.cn/base/'):
        url_get1 = loc_url.split('?')[0]
        params1 = loc_url.split('?')[1]
        app_token_tmp = url_get1.split('/')[-1]
        table_id_tmp = params1.split('=')[1].split('&')[0]
        logging.info(f'feishu-config.ini parser app_token:{app_token_tmp} table_id:{table_id_tmp}')
        configure.set('SECRET', 'app_token', app_token_tmp)
        configure.set('SECRET', 'table_id', table_id_tmp)
    else:
        return False

    at = configure.get('SECRET', 'app_token')
    ti = configure.get('SECRET', 'table_id')
    logging.debug(f'feishu-config.ini app_token:{at}, table_id:{ti}')

    cl_qrcode = configure.get('颜色', '二维码颜色')
    cl_text = configure.get('颜色', '文本颜色')

    # 读取配置文件中的说明
    desc = configure.get('配置说明', '说明', fallback=None)
    logging.debug(f'feishu-config.ini desc:{desc}')

    # 读取配置文件中的字段
    qr_d_str = configure.get('二维码内容字段', '字段', fallback=None)
    qr_d = ast.literal_eval(qr_d_str)
    logging.debug(f'feishu-config.ini content_colum:{qr_d}')

    # 读取配置文件中的文本显示
    left_str = configure.get('文本显示', '位置_左右', fallback=None)
    text_conf.left = left_str == "右"
    content_colum = configure.get('文本显示', '字段', fallback=None)
    text_conf.items = ast.literal_eval(content_colum)
    font_size_str = configure.get('文本显示', '字体大小', fallback=None)
    text_conf.font_size = int(font_size_str)
    logging.debug(f'feishu-config.ini text_conf:{str(text_conf)} {text_conf.font_size}')

    # 纸张配置
    label_row_col_str = configure.get('标签', '标签行列', fallback=None)
    label_row_col = ast.literal_eval(label_row_col_str)
    rows = label_row_col[0]
    cols = label_row_col[1]
    page_conf.row_colum = (rows, cols)

    paper_size_str = configure.get('打印纸张', 'page_size', fallback=None)
    page_conf.page_size = page_size = ast.literal_eval(paper_size_str)
    page_margin_str = configure.get('打印纸张', '边距_上下左右', fallback=None)
    page_conf.page_margin = page_margin = ast.literal_eval(page_margin_str)

    page_conf.content_size_cm = content_size_cm = (page_size[0] - page_margin[2] - page_margin[3],
                                                   page_size[1] - page_margin[0] - page_margin[1])
    page_conf.content_size_pixels = (cm_to_pixels(content_size_cm[0]), cm_to_pixels(content_size_cm[1]))

    page_conf.label_size_cm = label_size_cm = (content_size_cm[0] / cols, content_size_cm[1] / rows)
    page_conf.label_size_pixels = (cm_to_pixels(label_size_cm[0]), cm_to_pixels(label_size_cm[1]))

    page_conf.min_direction_pixels = min(min(page_conf.label_size_pixels[0], page_conf.label_size_pixels[1]), 500)

    # 扫码设置
    mobile_conf.url = loc_url
    mobile_conf.qr_code_data = qr_d
    editable_str = configure.get('扫码设置', '可编辑字段', fallback=None)
    editable = ast.literal_eval(editable_str)
    mobile_conf.editable = editable
    logging.debug(f'feishu-config.ini ed:{editable}  {type(editable)}')

    return True


def reset_config_file():
    if os.path.exists('./feishu-config.ini'):
        os.remove('./feishu-config.ini')
    configure.remove_section('SECRET')
    check_config_file()
    logging.info('完成重置配置')


def get_big_picture_draw() -> (Image, ImageDraw):
    new_image = Image.new("RGB", page_conf.content_size_pixels, bg_content)
    return new_image, ImageDraw.Draw(new_image)


def get_docx_document() -> Document:
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal'].paragraph_format.line_spacing = Pt(0)
    section = document.sections[0]
    section.page_width = Cm(page_conf.page_size[0])
    section.page_height = Cm(page_conf.page_size[1])
    section.top_margin = Cm(page_conf.page_margin[0])
    section.bottom_margin = Cm(page_conf.page_margin[1])
    section.left_margin = Cm(page_conf.page_margin[2])
    section.right_margin = Cm(page_conf.page_margin[3])
    return document


def save_big_picture_to_docx(document: Document, img: Image):
    base_img_file_path = f"./qrcodes/qrcode_base.png"
    img.save(base_img_file_path)
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(base_img_file_path, width=Cm(page_conf.content_size_cm[0]), height=Cm(page_conf.content_size_cm[1]))


def insert_images_into_docx():
    rows = page_conf.row_colum[0]
    cols = page_conf.row_colum[1]
    cell_index = 0
    folder_path = './qrcodes'
    document = get_docx_document()
    image_files = [f for f in os.listdir(folder_path) if f.endswith('.png')]
    images_per_page = rows * cols
    print(f" images_per_page {images_per_page} image_files size={len(image_files)}")
    img: Image = None
    draw: ImageDraw = None
    for i, image_file in enumerate(image_files, start=1):
        if i % images_per_page == 1:
            if img is not None:
                save_big_picture_to_docx(document, img)
            img, draw = get_big_picture_draw()
            cell_index = 0

        # 获取当前单元格
        cell_position = (cell_index // cols, cell_index % cols)
        print(f"cell_position = {cell_position}")
        image_path = os.path.join(folder_path, image_file)
        label = Image.open(image_path)
        label_position = (cell_position[1] * page_conf.label_size_pixels[0],
                          cell_position[0] * page_conf.label_size_pixels[1])
        label_margin_size = (page_conf.label_size_pixels[0]-page_conf.label_margin_pixels,
                             page_conf.label_size_pixels[1]-page_conf.label_margin_pixels)
        label = label.resize(label_margin_size)
        fix_position = (int(label_position[0]+page_conf.label_margin_pixels/2),
                        int(label_position[1]+page_conf.label_margin_pixels/2))
        img.paste(label, fix_position)
        cell_index += 1
    save_big_picture_to_docx(document, img)

    # 先清空目录下的所有文件
    for f in os.listdir('./qrcodes'):
        try:
            os.remove(os.path.join('./qrcodes', f))
        except Exception as e:
            logging.info(f'####got exception:{e}')
            continue
    # try 刪除目录./qrcodes
    try:
        os.rmdir('./qrcodes')
    except Exception as e:
        logging.info(f'####got exception:{e}')
        pass

    time_str = time.strftime("打印%Y%m%d%H%M%S", time.localtime())
    document.save(f'./{time_str}.docx')
    logging.info(f"当前文件夹 输出{time_str}.docx")
    os.system(f'start ./{time_str}.docx')

def get_tenant_access_token(app_id=None, app_secret=None, config_file=None):
    # 构建请求URL和请求头
    url_access = "https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal"
    headers = {
        'Content-Type': 'application/json; charset=utf-8',
    }

    # 构建请求体
    payload = {
        "app_id": app_id,
        "app_secret": app_secret
    }

    # 发起请求
    response = requests.post(url_access, headers=headers, data=json.dumps(payload))
    response_json = response.json()
    return response_json.get('code'), response_json.get('msg'), response_json.get('tenant_access_token')


def list_records(tenant_access_token_p=None, app_token_p=None, table_id_p=None, page_token=None, page_size=None,
                 config_file=None):
    url_record = f"https://open.feishu.cn/open-apis/bitable/v1/apps/{app_token_p}/tables/{table_id_p}/records?filter=CurrentValue.%5B%E5%8B%BE%E9%80%89%E6%89%93%E5%8D%B0%E6%A0%87%E7%AD%BE%5D%3D1"
    headers = {
        'Authorization': 'Bearer ' + tenant_access_token_p,
        'Content-Type': 'application/json; charset=utf-8',
    }
    params = {'page_size': page_size}
    if page_token:
        params['page_token'] = page_token

    response = requests.get(url_record, headers=headers, params=params)
    res = response.json()
    data = res.get('data', {})
    code = res.get('code', -1)
    items = data.get('items', [])
    page_token = data.get('page_token', '')
    has_more = data.get('has_more', False)
    total = data.get('total', 0)
    msg = res.get('msg')
    ##logging.info 改为logging
    logging.info(f'code:{code} msg:{msg} total:{total} has_more:{has_more} page_token:{page_token}')
    return code, msg, items


def get_simple_qr_data(item):
    qr_data_ = {}
    text_data_ = {}
    try:
        fields = item.get('fields', {})
        qr_data_['Uid'] = item.get('record_id', '')
        # qr_data['User'] = fields.get('使用人', [{}])[0].get('name', '')
        for key in ast.literal_eval(configure.get('二维码内容字段', '字段')):
            if '日期' in key or '时间' in key or 'time' in key:
                time_stamp = fields.get(key, 0)
                qr_data_[key] = time.strftime("%Y%m%d", time.localtime(time_stamp / 1000))
            else:
                qr_data_[key] = fields.get(key, '')

        for key in text_conf.items:
            if '日期' in key or '时间' in key or 'time' in key:
                time_stamp = fields.get(key, 0)
                text_data_[key] = time.strftime("%Y%m%d", time.localtime(time_stamp / 1000))
            else:
                text_data_[key] = fields.get(key, '')
        # logging.info(f'text_data_:{text_data_}')
    except Exception as e:
        logging.info(f'####got exception:{e}')
        qr_data_ = None
    return text_data_, qr_data_


def simplify_records(items_) -> list:
    for item in items_:
        text_data__, qr_data__ = get_simple_qr_data(item)
        if qr_data__ is None or text_data__ is None:
            continue
        yield text_data__, qr_data__


def gen_qrcode_by_qr_data(_text_data, _qr_data):
    # QR Code
    qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_L, box_size=10, border=4)
    qr_info = str()
    for key in _qr_data:
        qr_info = "".join([qr_info, key, ": ", _qr_data[key], '\n'])
    # logging.info(f'join _qr_data result {qr_info}')
    qr.add_data(qr_info)
    qr.make(fit=True)
    img = qr.make_image(fill_color=cl_qrcode, back_color=bg_qrcode)
    img = img.resize((page_conf.min_direction_pixels, page_conf.min_direction_pixels))
    # img.save(f"./qrcodes/PPPPPPP.png")

    # TEXT
    text_info = str()
    for key in _text_data:
        text_info = "".join([text_info, key, ": ", _text_data[key], '\n'])
    text_lines = text_info.split('\n')
    font_size = text_conf.font_size
    font = ImageFont.truetype("msyhl.ttc", font_size)

    # left or not
    is_left = text_conf.left

    new_image = Image.new("RGB", (page_conf.label_size_pixels[0], page_conf.label_size_pixels[1]), bg_label)
    draw = ImageDraw.Draw(new_image)
    if is_left:
        text_position = (20, 25)
        # Draw text
        for line in text_lines:
            if "Uid" in line:
                continue
            draw.text(text_position, line, font=font, fill=cl_text)
            text_position = (text_position[0], text_position[1] + font_size + 10)
        # Draw Qrcode
        new_image.paste(img, (page_conf.label_size_pixels[0] - page_conf.min_direction_pixels, 0))
    else:
        # Draw text
        text_position = (page_conf.min_direction_pixels + 10, 25)
        for line in text_lines:
            if "Uid" in line:
                continue
            draw.text(text_position, line, font=font, fill=cl_text)
            text_position = (text_position[0], text_position[1] + font_size + 10)

        # Draw Qrcode
        new_image.paste(img, (0, 0))

    # Save the result
    new_image.save(f"./qrcodes/qrcode_{_qr_data['Uid']}.png")


def print_log(log_str):
    print(log_str)
    logging.info(log_str)
    pass


def gen_mobile_qrcode():
    check_config_file()
    mobile_qrcode_dict = dict()
    mobile_qrcode_dict['url'] = mobile_conf.url
    mobile_qrcode_dict['item'] = list(mobile_conf.qr_code_data)
    mobile_qrcode_dict['editable'] = list(mobile_conf.editable)
    setting_js = json.dumps(mobile_qrcode_dict, ensure_ascii=False)
    qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_L, box_size=10, border=4)
    qr.add_data(setting_js)
    qr.make(fit=True)
    img = qr.make_image(fill_color=cl_qrcode, back_color=bg_qrcode)
    p = f"./mobile_settings_qrcode.png"
    img.save(p)
    os.system(f'start {p}')
    pass


def do_main():
    """如果文件qrcodes目录不存在，则创建"""
    if not os.path.exists('./qrcodes'):
        os.makedirs('./qrcodes')
    check_config_file()
    code1, msg1, tenant_access_token = get_tenant_access_token(app_id=app_id, app_secret=app_secret)
    # logging.info(f'code1:{code1} tenant_access_token:{tenant_access_token}')

    app_token_cf = configure.get('SECRET', 'app_token')
    table_id_cf = configure.get('SECRET', 'table_id')
    if code1 == 0:
        code2, msg2, items = list_records(tenant_access_token_p=tenant_access_token, app_token_p=app_token_cf,
                                          table_id_p=table_id_cf)
        logging.info(f'获取到标签记录数量:{len(items)}')
        i = 0
        if code2 == 0:
            for text_data, qr_data in simplify_records(items):
                i = i+1
                logging.info(f'第{i}个标签\n标签文本内容:{text_data}\n标签二维码内容:{qr_data}')
                gen_qrcode_by_qr_data(text_data, qr_data)
            insert_images_into_docx()
            logging.info(f'导出标签打印文档成功!')
        else:
            logging.info(f'获取表格记录失败!!, code:{code2}')
    else:
        logging.info(f'被拒绝访问多维表格，请确认【多维表格权限配置是否正确】和【是否在允许的网络域内】!!, code:{code1}')

# do_main()