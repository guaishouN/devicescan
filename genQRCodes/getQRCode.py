import requests
import configparser
import json
import qrcode
from PIL import Image, ImageDraw, ImageFont
import time
from docx import Document
from docx.shared import Cm
import logging
import os

default_table_url = "https://yesv-desaysv.feishu.cn/base/CmHmb4MxPaEW7zsWB07c1hCUnhd?table=tbl3OBzMMqjX79gN&view=vewsIt61jC#CategoryScheduledTask"
app_id = 'cli_a514aea9fa79900b'
app_secret = 'IsUeIxmzO5NtJiQA6B3MdfkHqIcmQqws'
app_token = 'CmHmb4MxPaEW7zsWB07c1hCUnhd'
table_id = 'tbl3OBzMMqjX79gN'


"""如果文件qrcodes目录不存在，则创建"""
if not os.path.exists('./qrcodes'):
    os.makedirs('./qrcodes')

"""log 以追加的形式打印到文件夹./qrcodes下"""
# logging.basicConfig(level=logging.DEBUG, filename='./qrcodes/temp_qrlog.log', filemode='a', format='%(asctime)s - %(
# levelname)s - %(message)s')
logging.basicConfig(level=logging.DEBUG, filemode='a', format='%(asctime)s - %(levelname)s - %(message)s')
configure = configparser.ConfigParser()


def check_config_file():
    global configure
    # 解析并读取配置文件feishu-config.ini
    configure["多维表格地址"] = {'url': default_table_url}
    # 在config文件中加入注释"单位为cm"
    desc = "#以下为默认配置。单位为cm, A4纸打印，大小为21cm*29.7cm。"
    configure["配置说明"] = {"说明": desc}
    configure["打印纸张"] = {'边距_上下左右': (1.56, 1.3, 0.78, 0.78)}
    configure["标签"] = {'标签行列': (7, 3), '大小': (6.3, 3.8), '标签边距_上下左右': (0, 0, 0, 0.2)}
    # 检查飞书配置文件  feishu-config.ini，如果不存在，则创建
    if not os.path.exists('./feishu-config.ini'):
        # utf-8 编码写入配置文件
        with open('./feishu-config.ini', 'w', encoding='utf-8') as cf:
            configure.write(cf)
    else:
        configure.read('feishu-config.ini', encoding='utf-8')
    paper_size = (21.0, 29.7)
    configure.set("打印纸张", 'page_size', str(paper_size))
    configure.add_section('SECRET')
    configure.set('SECRET', 'app_id', app_id)
    configure.set('SECRET', 'app_secret', app_secret)
    configure.set('SECRET', 'app_token', app_token)
    configure.set('SECRET', 'table_id', table_id)

    # 读取配置文件中的url, 如果没有，则使用默认值
    loc_url = configure.get('多维表格地址', 'url', fallback=None)
    logging.info(f'feishu-config.ini url:{loc_url}')
    if not loc_url:
        loc_url = default_table_url
        logging.info(f'feishu-config.ini default url:{loc_url}')

    # 如果url符合格式则从url中解析出app_token table_id
    if loc_url.startswith('https://yesv-desaysv.feishu.cn/base/'):
        url_get1 = loc_url.split('?')[0]
        params1 = loc_url.split('?')[1]
        app_token_tmp = url_get1.split('/')[-1]
        table_id_tmp = params1.split('=')[1].split('&')[0]
        logging.info(f'feishu-config.ini parser app_token:{app_token_tmp} table_id:{table_id_tmp}')
        configure.set('SECRET', 'app_token', app_token_tmp)
        configure.set('SECRET', 'table_id', table_id_tmp)
    else:
        return False

    at = configure.get('SECRET', 'app_token')
    ti = configure.get('SECRET', 'table_id')
    logging.info(f'feishu-config.ini app_token:{at}, table_id:{ti}')

    # 读取配置文件中的说明
    desc = configure.get('配置说明', '说明', fallback=None)
    logging.info(f'feishu-config.ini desc:{desc}')
    # 读取配置文件中的纸张大小, paper_size = configure.get('打印纸张', '宽高', fallback=None)
    paper_size = configure.get('打印纸张', 'page_size', fallback=None)
    logging.info(f'feishu-config.ini paper_size:{paper_size}')
    # 读取配置文件中的边距
    paper_margin = configure.get('打印纸张', '边距_上下左右', fallback=None)
    logging.info(f'feishu-config.ini paper_margin:{paper_margin}')
    # 读取配置文件中的标签行列
    label_row_col = configure.get('标签', '标签行列', fallback=None)
    logging.info(f'feishu-config.ini label_row_col:{label_row_col}')
    # 读取配置文件中的标签大小
    label_size = configure.get('标签', '大小', fallback=None)
    logging.info(f'feishu-config.ini label_size:{label_size}')
    # 读取配置文件中的标签大小
    label_margin = configure.get('标签', '标签边距_上下左右', fallback=None)
    logging.info(f'feishu-config.ini label_margin:{label_margin}')
    return True


def get_tenant_access_token(app_id=None, app_secret=None, config_file=None):
    # 构建请求URL和请求头
    url_access = "https://open.feishu.cn/open-apis/auth/v3/tenant_access_token/internal"
    headers = {
        'Content-Type': 'application/json; charset=utf-8',
    }

    # 构建请求体
    payload = {
        "app_id": app_id,
        "app_secret": app_secret
    }

    # 发起请求
    response = requests.post(url_access, headers=headers, data=json.dumps(payload))
    response_json = response.json()
    return response_json.get('code'), response_json.get('msg'), response_json.get('tenant_access_token')


def list_records(tenant_access_token_p=None, app_token_p=None, table_id_p=None, page_token=None, page_size=None,
                 config_file=None):
    url_record = f"https://open.feishu.cn/open-apis/bitable/v1/apps/{app_token_p}/tables/{table_id_p}/records?filter=CurrentValue.%5B%E5%8B%BE%E9%80%89%E6%89%93%E5%8D%B0%E6%A0%87%E7%AD%BE%5D%3D1"
    headers = {
        'Authorization': 'Bearer ' + tenant_access_token_p,
        'Content-Type': 'application/json; charset=utf-8',
    }
    params = {'page_size': page_size}
    if page_token:
        params['page_token'] = page_token

    response = requests.get(url_record, headers=headers, params=params)
    res = response.json()
    data = res.get('data', {})
    code = res.get('code', -1)
    items = data.get('items', [])
    page_token = data.get('page_token', '')
    has_more = data.get('has_more', False)
    total = data.get('total', 0)
    msg = res.get('msg')
    ##logging.info 改为logging
    logging.info(f'code:{code} msg:{msg} total:{total} has_more:{has_more} page_token:{page_token}')
    return code, msg, items


def get_simple_qr_data(item):
    # item to qr_data
    try:
        qr_data = {}
        fields = item.get('fields', {})
        qr_data['Uid'] = item.get('record_id', '')
        # qr_data['User'] = fields.get('使用人', [{}])[0].get('name', '')
        qr_data['User'] = fields.get('使用人', '')
        # date format: 2024/01/05
        time_stamp = fields.get('录入日期', 0)
        qr_data['Date'] = time.strftime("%Y%m%d", time.localtime(time_stamp / 1000))
        qr_data['Dev'] = fields.get('设备名称', '')
        qr_data['DevId'] = fields.get('设备ID', '')
        qr_data['Proj'] = fields.get('项目', '')
        # logging.info(f'qr_data:{qr_data}')
    except Exception as e:
        logging.info(f'####got exception:{e}')
        qr_data = None
    return qr_data


def simplify_records(items) -> list:
    for item in items:
        qr_data = get_simple_qr_data(item)
        if qr_data is None:
            continue
        yield qr_data


def gen_qrcode_by_qr_data(qr_data):
    qr = qrcode.QRCode(version=1, error_correction=qrcode.constants.ERROR_CORRECT_L, box_size=10, border=4)
    info = f"Uid：{qr_data['Uid']}\nDevId：{qr_data['DevId']}\nDev：{qr_data['Dev']}\nProj：{qr_data['Proj']}\nUser：{qr_data['User']}\nDate：{qr_data['Date']}"
    # logging.info(f'{info}')
    qr.add_data(info)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    # img.save(f"./qrcodes/PPPPPPP.png")
    text_lines = info.split('\n')
    # Create a new image with white background
    new_image = Image.new("RGB", (400 + 480, 480), "white")

    # Get a drawing context
    draw = ImageDraw.Draw(new_image)

    # Set a font
    # font = ImageFont.load_default()
    font_size = 34
    font = ImageFont.truetype("msyhl.ttc", font_size)

    # Set the position to start drawing the text
    text_position = (20, 50)

    # Write each line of text
    for line in text_lines:
        if "Uid" in line:
            continue
        draw.text(text_position, line, font=font, fill="black")
        text_position = (text_position[0], text_position[1] + font_size + 18)

    # Paste the original image on the right side
    new_image.paste(img, (400, 0))

    # Save the result
    new_image.save(f"./qrcodes/qrcode_{qr_data['Uid']}.png")


def insert_images_into_word():
    folder_path = './qrcodes'
    # 创建一个Word文档对象，A4纸张 210mm*297mm
    document = Document()
    # 设置页面大小为A4纸张，210mm*297mm
    document.styles['Normal'].font.name = u'宋体'
    # 行间距为0倍字体大小
    document.styles['Normal'].paragraph_format.line_spacing = 1.11

    # 设置A4纸张，页距：上下各1.54cm 左右各0.78cm，单标签大小：宽6.3cm 高3.8
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.56)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(0.78)
    section.right_margin = Cm(0.78)
    # 遍历文件夹中的图片文件
    image_files = [f for f in os.listdir(folder_path) if f.endswith(('.png'))]

    # 每页3列*7行的方式排版
    rows = 7
    cols = 3
    images_per_page = rows * cols

    for i, image_file in enumerate(image_files, start=1):
        # 在新的一页开始时，添加一个表格, 每页3列*7行的方式排版
        # 表格固定width=6.3cm height=3.7cm     
        if i % images_per_page == 1:
            table = document.add_table(rows=rows, cols=cols)
            for row in table.rows:
                row.height = Cm(3.7)
            for col in table.columns:
                col.width = Cm(6.3)
            # table.style = 'Table Grid'
            cell_index = 0

        # 获取当前单元格
        cell = table.cell(cell_index // cols, cell_index % cols)

        # 插入图片到表格单元格
        image_path = os.path.join(folder_path, image_file)
        # 顶格插入图片        
        # cell.add_paragraph().add_run().add_picture(image_path, width=Cm(6.3), height=Cm(3.8))

        # 居中插入图片
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_path, width=Cm(6.2), height=Cm(3.5))
        run.alignment = 2
        cell_index += 1

    # 先清空目录下的所有文件
    for f in os.listdir('./qrcodes'):
        try:
            os.remove(os.path.join('./qrcodes', f))
        except Exception as e:
            logging.info(f'####got exception:{e}')
            continue
    # try 刪除目录./qrcodes
    try:
        os.rmdir('./qrcodes')
    except Exception as e:
        logging.info(f'####got exception:{e}')
        pass

    # 保存文档
    time_str = time.strftime("打印%Y%m%d%H%M%S", time.localtime())
    document.save(f'./{time_str}.docx')


def print_log(log_str):
    print(log_str)
    logging.info(log_str)
    pass


check_config_file()
code1, msg1, tenant_access_token = get_tenant_access_token(app_id=app_id, app_secret=app_secret)
logging.info(f'code1:{code1} tenant_access_token:{tenant_access_token}')

app_token_cf = configure.get('SECRET', 'app_token')
table_id_cf = configure.get('SECRET', 'table_id')
if code1 == 0:
    code2, msg2, items = list_records(tenant_access_token_p=tenant_access_token, app_token_p=app_token_cf,
                                      table_id_p=table_id_cf)
    logging.info(f'code2:{code2} records sizes:{len(items)}')
    if code2 == 0:
        for qr_data in simplify_records(items):
            logging.info(f'qr_data:{qr_data["DevId"]}')
            gen_qrcode_by_qr_data(qr_data)
        insert_images_into_word()
        logging.info(f'Suceessfully generated qrcodes and inserted into word, done!')
    else:
        logging.info(f'list_records failed!!, code:{code2}')
else:
    logging.info(f'get_tenant_access_token failed!!, code:{code1}')
